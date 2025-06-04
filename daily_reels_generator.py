import os
import random
import shutil
from pathlib import Path
from moviepy.editor import VideoFileClip
from docx import Document

VIDEOS_DIR = Path("videos")
OUTPUT_DIR = Path("final_reels")
LOG_FILE = Path("Published_Videos_Log.docx")

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# 📄 تحميل سجل الفيديوهات
if LOG_FILE.exists():
    doc = Document(LOG_FILE)
else:
    doc = Document()
    doc.add_heading("Published Videos Log", 0)

used_videos = set(p.text for p in doc.paragraphs[1:] if p.text)

# 🔍 العثور على فيديوهات جديدة تحتوي على صوت فقط وتكون رأسية بدقة جيدة
available_videos = []
for subfolder in VIDEOS_DIR.glob("*/"):
    for video_file in subfolder.glob("*.mp4"):
        print(f"\n🔍 فحص: {video_file.name}")

        if video_file.name in used_videos:
            print("⏭️ تم استخدامه من قبل")
            continue

        try:
            clip = VideoFileClip(str(video_file))
            w, h = clip.size
            duration = clip.duration
            audio = clip.audio

            if w >= h:
                print("❌ مستبعد: الفيديو أفقي")
                continue
            if clip.audio is None:
                print("❌ مستبعد: لا يحتوي على صوت")
                continue
            if h < 720:
                print("❌ مستبعد: الجودة أقل من 720p")
                continue

            print("✅ مؤهل: سيتم إضافته")
            available_videos.append(video_file)
        except Exception as e:
            print(f"⚠️ خطأ عند فحص {video_file.name}: {e}")

# 🎯 اختيار 2 فيديو عشوائي فقط
selected = random.sample(available_videos, k=min(2, len(available_videos)))

for video_path in selected:
    dest_path = OUTPUT_DIR / video_path.name
    shutil.copy(video_path, dest_path)
    print(f"📁 تم النسخ إلى final_reels/: {video_path.name}")
    doc.add_paragraph(video_path.name)

# 💾 حفظ السجل
doc.save(LOG_FILE)
print("📄 تم تحديث سجل الفيديوهات.")


import os
import json
import tempfile
from pathlib import Path
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document
from datetime import datetime

OUTPUT_DIR = "output_reels"

if not os.path.exists(OUTPUT_DIR):
    print(f"⚠️ المجلد {OUTPUT_DIR} غير موجود، لا يوجد شيء لرفعه.")
    exit(0)

service_account_info = json.loads(os.environ["SERVICE_ACCOUNT_JSON"])

with tempfile.NamedTemporaryFile(mode="w+", delete=False, suffix=".json") as f:
    json.dump(service_account_info, f)
    SERVICE_ACCOUNT_FILE = f.name

SCOPES = ['https://www.googleapis.com/auth/drive']
credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE, scopes=SCOPES)
drive_service = build('drive', 'v3', credentials=credentials)

def log_uploaded_video(keyword, filename):
    log_file = "Published_Videos_Log.docx"
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    if os.path.exists(log_file):
        doc = Document(log_file)
    else:
        doc = Document()
        doc.add_heading("Published Videos Log", 0)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Keyword'
        hdr_cells[2].text = 'Filename'

    table = doc.tables[0]
    row_cells = table.add_row().cells
    row_cells[0].text = now
    row_cells[1].text = keyword
    row_cells[2].text = filename

    doc.save(log_file)

def upload_file(file_path, parent_id, keyword):
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [parent_id]
    }
    media = MediaFileUpload(file_path, resumable=True)
    uploaded = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f"✅ تم رفع الملف: {file_path}")
    log_uploaded_video(keyword, os.path.basename(file_path))

def get_or_create_folder(name, parent_id=None):
    query = f"name='{name}' and mimeType='application/vnd.google-apps.folder'"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
    folders = results.get('files', [])
    if folders:
        return folders[0]['id']
    file_metadata = {'name': name, 'mimeType': 'application/vnd.google-apps.folder'}
    if parent_id:
        file_metadata['parents'] = [parent_id]
    folder = drive_service.files().create(body=file_metadata, fields='id').execute()
    return folder.get('id')

def upload_output_videos():
    root_id = get_or_create_folder("AutoUploader")
    final_videos_id = get_or_create_folder("final_videos", parent_id=root_id)
    for keyword_folder in Path(OUTPUT_DIR).iterdir():
        if keyword_folder.is_dir():
            keyword = keyword_folder.name
            keyword_drive_id = get_or_create_folder(keyword, parent_id=final_videos_id)
            for video_file in keyword_folder.glob("*.mp4"):
                upload_file(str(video_file), parent_id=keyword_drive_id, keyword=keyword)

if __name__ == "__main__":
    upload_output_videos()

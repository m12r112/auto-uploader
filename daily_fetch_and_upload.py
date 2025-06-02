import os
import json
import requests
import random
from bs4 import BeautifulSoup
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document

# إعداد التفويض باستخدام GitHub Secret
SCOPES = ['https://www.googleapis.com/auth/drive']
key_data = json.loads(os.environ["SERVICE_ACCOUNT_KEY"])
credentials = service_account.Credentials.from_service_account_info(key_data, scopes=SCOPES)
service = build('drive', 'v3', credentials=credentials)

# إعدادات المجلدات
with open("settings.json") as f:
    settings = json.load(f)

TEMP_FOLDER = settings["local_temp_folder"]
DRIVE_MAIN_FOLDER = settings["main_drive_folder"]

# تحميل الكلمات المفتاحية
with open("keywords.json") as f:
    keywords = json.load(f)

# تحقق من وجود ملف السجل، وإنشاؤه إذا لم يكن موجودًا
if not os.path.exists("Published_Videos_Log.docx"):
    doc = Document()
    doc.add_paragraph("سجل الفيديوهات المنشورة")
    doc.save("Published_Videos_Log.docx")

# سجل الفيديوهات
published_doc = Document("Published_Videos_Log.docx")
published_titles = [p.text.strip() for p in published_doc.paragraphs if p.text.strip()]

def fetch_video_links(keyword):
    url = f"https://www.pexels.com/search/videos/{keyword}/"
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        video_tags = soup.find_all("video")
        links = [tag.get("src") for tag in video_tags if tag.get("src")]
        return links
    except Exception as e:
        print(f"⚠️ خطأ في جلب الفيديوهات لـ {keyword}: {e}")
        return []

def upload_to_drive(filepath, filename, parent_folder_id):
    file_metadata = {'name': filename, 'parents': [parent_folder_id]}
    media = MediaFileUpload(filepath, resumable=True)
    uploaded_file = service.files().create(
        body=file_metadata, media_body=media, fields='id').execute()
    return uploaded_file.get('id')

for kw in keywords:
    folder_path = os.path.join(TEMP_FOLDER, kw)
    os.makedirs(folder_path, exist_ok=True)

    print(f"🔍 البحث عن: {kw}")
    video_links = fetch_video_links(kw)
    if not video_links:
        print(f"❌ لا توجد فيديوهات")
        continue

    selected = None
    for link in video_links:
        if link not in published_titles:
            selected = link
            break
    if not selected:
        print(f"⚠️ فيديوهات مكررة")
        continue

    filename = f"{kw}_{random.randint(1000,9999)}.mp4"
    file_path = os.path.join(folder_path, filename)

    try:
        with open(file_path, 'wb') as f:
            f.write(requests.get(selected, timeout=30).content)

        uploaded_id = upload_to_drive(file_path, filename, DRIVE_MAIN_FOLDER)
        print(f"✅ تم الرفع (ID: {uploaded_id})")

        published_doc.add_paragraph(selected)
        published_doc.save("Published_Videos_Log.docx")

        os.remove(file_path)  # حذف الملف لتوفير المساحة
    except Exception as e:
        print(f"❌ خطأ أثناء التحميل أو الرفع: {e}")

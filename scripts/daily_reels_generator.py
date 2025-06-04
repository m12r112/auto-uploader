import os
import random
from pathlib import Path
from moviepy.editor import VideoFileClip, AudioFileClip, concatenate_audioclips
from shutil import copy2
from docx import Document

# 🗂️ إعداد المسارات
VIDEOS_DIR = Path("videos")
AUDIO_DIR = Path("audio_library")
OUTPUT_DIR = Path("final_reels")
LOG_FILE = Path("Published_Videos_Log.docx")

# 🔑 تحميل التوكن من المتغير البيئي
access_token = os.environ.get("INSTAGRAM_ACCESS_TOKEN", "")
if not access_token:
    print("⚠️ لم يتم العثور على توكن Instagram. سيتم تخطي النشر.")

# 🧠 تحميل السجل الحالي لتجنب التكرار
if LOG_FILE.exists():
    doc = Document(LOG_FILE)
    published = [p.text for p in doc.paragraphs]
else:
    doc = Document()
    published = []

# 🧪 اختيار فيديوهين فقط من مجلدات مختلفة
all_videos = list(VIDEOS_DIR.rglob("*.mp4"))
random.shuffle(all_videos)
selected = []

for video_path in all_videos:
    if len(selected) >= 2:
        break
    if str(video_path) in published:
        continue

    try:
        clip = VideoFileClip(str(video_path))
        if clip.w < clip.h and clip.duration <= 60 and clip.fps >= 24:
            selected.append((video_path, clip))
    except Exception as e:
        print(f"❌ خطأ في فتح الفيديو {video_path}: {e}")

# 🧩 معالجة الفيديوهات
for video_path, clip in selected:
    keyword = video_path.parent.name
    has_audio = clip.audio is not None

    output_path = OUTPUT_DIR / video_path.name
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if has_audio:
        # ✅ يحتوي على صوت: يُنسخ فقط
        copy2(video_path, output_path)
        print(f"✅ تم نسخ الفيديو {video_path.name} كما هو.")
    else:
        # 🔇 لا يحتوي على صوت: دمج مع صوت مناسب
        audio_folder = AUDIO_DIR / keyword
        audio_files = list(audio_folder.glob("*.mp3"))

        if not audio_files:
            print(f"⚠️ لا يوجد صوت مناسب لـ {keyword}")
            continue

        used_audio = audio_files.pop(0)
        audio_clip = AudioFileClip(str(used_audio))

        if audio_clip.duration > clip.duration:
            audio_clip = audio_clip.subclip(0, clip.duration)
        else:
            loops = int(clip.duration // audio_clip.duration) + 1
            audio_clip = concatenate_audioclips([audio_clip] * loops).subclip(0, clip.duration)

        final = clip.set_audio(audio_clip)
        final.write_videofile(str(output_path), codec="libx264", audio_codec="aac")
        print(f"🎧 تم دمج صوت مع الفيديو {video_path.name}")

    # ✍️ تسجيل الفيديو في السجل
    doc.add_paragraph(str(video_path))
    print(f"📝 تم تسجيل الفيديو {video_path.name} في السجل.")

# 💾 حفظ ملف السجل
doc.save(LOG_FILE)
print("✅ تم حفظ السجل.")

# 📤 (اختياري) نشر الفيديوهات باستخدام التوكن
if access_token:
    print("📤 يمكنك الآن استخدام التوكن للنشر على Instagram (لم يُنفّذ هنا).")

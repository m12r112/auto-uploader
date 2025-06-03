import smtplib
import os
import json
from datetime import datetime, timedelta
from email.message import EmailMessage
from docx import Document

LOG_FILE = "Published_Videos_Log.docx"
TO_EMAIL = "mm2677078@gmail.com"
FROM_EMAIL = "mm2677078@gmail.com"
APP_PASSWORD = "mijcyxkkyhyaihwx"

def extract_weekly_entries(log_file):
    if not os.path.exists(log_file):
        return []

    doc = Document(log_file)
    table = doc.tables[0]
    one_week_ago = datetime.now() - timedelta(days=7)
    rows = []

    for row in table.rows[1:]:
        date_str = row.cells[0].text.strip()
        try:
            date_obj = datetime.strptime(date_str, "%Y-%m-%d %H:%M")
            if date_obj >= one_week_ago:
                keyword = row.cells[1].text.strip()
                filename = row.cells[2].text.strip()
                rows.append((date_str, keyword, filename))
        except:
            continue

    return rows

def create_report(entries, filename):
    doc = Document()
    doc.add_heading("Weekly AutoUploader Report", 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Keyword'
    hdr_cells[2].text = 'Filename'
    for date, keyword, filename in entries:
        row = table.add_row().cells
        row[0].text = date
        row[1].text = keyword
        row[2].text = filename
    doc.save(filename)

def send_email_report():
    entries = extract_weekly_entries(LOG_FILE)
    if not entries:
        print("📭 لا توجد بيانات كافية لإرسال التقرير الأسبوعي.")
        return

    today = datetime.now().strftime("%Y-%m-%d")
    report_name = f"Weekly_Report_{today}.docx"
    create_report(entries, report_name)

    msg = EmailMessage()
    msg["Subject"] = f"📊 Weekly AutoUploader Report – {today}"
    msg["From"] = FROM_EMAIL
    msg["To"] = TO_EMAIL
    msg.set_content("""Hello Mohammed,

Attached is your weekly AutoUploader report showing all videos uploaded in the past 7 days.

Best regards,
AutoUploader Bot 🤖
""")

    with open(report_name, "rb") as f:
        file_data = f.read()
        msg.add_attachment(file_data, maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=report_name)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(FROM_EMAIL, APP_PASSWORD)
        smtp.send_message(msg)

    print(f"📨 تم إرسال التقرير الأسبوعي إلى {TO_EMAIL} بنجاح!")

if __name__ == "__main__":
    send_email_report()
